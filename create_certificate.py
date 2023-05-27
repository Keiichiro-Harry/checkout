import os
import pandas as pd
from docx import Document

# テンプレートファイルのパス
template_file = 'format.docx'
# 保存先フォルダ
save_folder = '無人航空機講習修了証明書'

# エクセルファイルからデータを読み込む
data = pd.read_excel('file.xlsx', sheet_name=1)

# データごとにワードファイルを作成
for _, row in data.iterrows():
    num_reg = row['修了証明書番号']
    date_pub = row['発行日'].strftime('%Y年%m月%d日')
    date_upd = row['更新日'].strftime('%Y年%m月%d日')
    name_fam = row['名前(姓)']
    name_fir = row['名前(名)']
    num_tec = str(row['技能証明書番号'])
    bool_cls1 = '〇' if row['初学1'] or row['経験1'] else ''
    bool_mlt1 = '〇' if row['初学1'] or row['経験1'] else ''
    bool_cls2 = '〇' if row['初学2'] or row['経験2'] else ''
    bool_mlt2 = '〇' if row['初学2'] or row['経験2'] else ''
    bool_dt1 = '〇' if row['目視外'] and bool_cls1 else ''
    bool_dt2 = '〇' if row['目視外'] and not bool_cls1 else ''
    bool_vis1 = '〇' if row['夜間'] and bool_cls1 else ''
    bool_vis2 = '〇' if row['夜間'] and not bool_cls1 else ''

    # テンプレートをコピーして新しいドキュメントを作成
    doc = Document(template_file)

    # プレースホルダーを置換する関数
    def replace_placeholder(doc, placeholder, value):
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, value)

    # プレースホルダーを置換する
    placeholders = {
        'NUM_REG': num_reg,
        'DATE_PUB': date_pub,
        'DATE_UPD': date_upd,
        'NAME_FAM': name_fam,
        'NAME_FIR': name_fir,
        'NUM_TEC': num_tec,
        'BOOL_CLS1': bool_cls1,
        'BOOL_CLS2': bool_cls2,
        'BOOL_MLT1': bool_mlt1,
        'BOOL_MLT2': bool_mlt2,
        'BOOL_DT1': bool_dt1,
        'BOOL_DT2': bool_dt2,
        'BOOL_VIS1': bool_vis1,
        'BOOL_VIS2': bool_vis2
    }

    for key, value in placeholders.items():
        replace_placeholder(doc, key, value)

    # 保存先フォルダが存在しない場合は作成する
    if not os.path.exists(save_folder):
        os.makedirs(save_folder)

    # ワードファイルを保存する
    file_name = f'{num_reg}_{name_fam}_{name_fir}.docx'
    file_path = os.path.join(save_folder, file_name)
    doc.save(file_path)
"""
import pandas as pd
from docxtpl import DocxTemplate
import os

# テンプレートファイルのパス
template_file = 'format.docx'

# エクセルファイルからデータを読み込む
data = pd.read_excel('file.xlsx', sheet_name=1)
     
# データごとにワードファイルを作成
for _, row in data.iterrows():
    # テンプレートをコピーして新しいドキュメントを作成
    doc = DocxTemplate(template_file)

    # 顧客情報を取得
    num_reg = row['修了証明書番号']
    date_pub = row['発行日'].strftime('%Y年%m月%d日')
    date_upd = row['更新日'].strftime('%Y年%m月%d日')
    name_fam = row['名前(姓)']
    name_fir = row['名前(名)']
    num_tec = str(row['技能証明書番号'])
    bool_cls1, bool_mlt1 = '〇', '〇' if row['初学1'] or row['経験1'] else ''
    bool_cls2, bool_mlt2 = '〇', '〇' if row['初学2'] or row['経験2'] else ''
    bool_dt1 = '〇' if row['目視外'] and bool_cls1 else ''
    bool_dt2 = '〇' if row['目視外'] and not bool_cls1 else ''
    bool_vis1 = '〇' if row['夜間'] and bool_cls1 else ''
    bool_vis2 = '〇' if row['夜間'] and not bool_cls1 else ''

    # テンプレート内の指定した場所に顧客の情報を挿入する
    placeholders = {
        'NUM_REG': num_reg,
        'DATE_PUB': date_pub,
        'DATE_UPD': date_upd,
        'NAME_FAM': name_fam,
        'NAME_FIR': name_fir,
        'NUM_TEC': num_tec,
        'BOOL_CLS1': bool_cls1,
        'BOOL_CLS2': bool_cls2,
        'BOOL_MLT1': bool_mlt1,
        'BOOL_MLT2': bool_mlt2,
        'BOOL_DT1': bool_dt1,
        'BOOL_DT2': bool_dt2,
        'BOOL_VIS1': bool_vis1,
        'BOOL_VIS2': bool_vis2
    }
    for paragraph in doc.paragraphs:
        for key, value in placeholders.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    
    doc.render(placeholders)
    # 保存先フォルダが存在しない場合は作成する
    save_folder = "無人航空機構講習修了証明書"
    if not os.path.exists(save_folder):
        os.makedirs(save_folder)

    # ワードファイルを保存する
    file_name = (f'{num_reg}_{name_fam}_{name_fir}.docx')
    file_path = os.path.join(save_folder, file_name)
    doc.save(file_path)
"""